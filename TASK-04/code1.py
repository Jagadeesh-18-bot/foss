from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, CallbackQueryHandler, ContextTypes
from docx import Document
import requests
import pandas as pd
import os

TELEGRAM_BOT_KEY=""
GOOGLE_BOOKS_API=""
BOT_USERNAME="technophiles-bot"

#start command
async def start_command(update: Update, context:ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Hello!, Im technophile. Please use the /help command to get started.")

#book command
async def book_command(update: Update, context:ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Please type in the genre along with the /book command in the format '/book <genre>'.")
        return
    genre = "+".join(context.args)
    csv_file = get_by_genre(genre)
    if csv_file:
        await update.message.reply_text("Here you go:")
        await update.message.reply_document(document=open(csv_file,'rb'),filename=f"{genre}_books.csv")
        os.remove(csv_file)
    else:
        await update.message.reply_text("Sorry, I couldn't find any books for this genre.")

def get_by_genre(genre):
    query = f'subject:{genre}'
    url = f'https://www.googleapis.com/books/v1/volumes?q={query}&maxResults=15&key={GOOGLE_BOOKS_API}'

    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        if 'items' in data:
            books = []
            for item in data['items']:
                volume_info = item['volumeInfo']
                book = {
                    "Title": volume_info.get("title", "Not available"),
                    "Author(s)": ", ".join(volume_info.get("authors", ["Not available"])),
                    "Published Date": volume_info.get("publishedDate", "Not available"),
                    "Description": volume_info.get("description", "Not available"),
                    "Preview Link": volume_info.get("previewLink", "Not available")
                }
                books.append(book)
            
            df = pd.DataFrame(books)
            csv_file = f"/tmp/{genre}_books.csv"
            df.to_csv(csv_file, index=False)
            
            return csv_file
        else:
            return None
    else:
        return "Error fetching data from Google Books API."

#preview command
async def preview_command(update: Update, context:ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Please type in the name of the book along with the /preview command in the format '/preview <book name>'.")
        return
    book_name = " ".join(context.args)
    preview_link = get_book_preview(book_name)
    await update.message.reply_text("Here's what I found:")
    await update.message.reply_text(preview_link)

def get_book_preview(book_name):
    query = book_name.replace(" ", "+")
    url = f'https://www.googleapis.com/books/v1/volumes?q={query}&key={GOOGLE_BOOKS_API}'

    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        if 'items' in data:
            book = data['items'][0]['volumeInfo']
            preview_link = book.get('previewLink', 'No preview link available')
            return preview_link
        else:
            return "No books found for the given title."
    else:
        return "Error fetching data from Google Books API."

#list command
async def list_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.args:
        book_name=" ".join(context.args)
        context.user_data['book_name']=book_name
        await update.message.reply_text("Got it!. Now use /reading_list command to add the book to your reading list or to remove the book from your reading list.")
    else:
        await update.message.reply_text("Please provide a book name along with the /list command in the format '/list <book name>' to proceed")

async def reading_list_command(update: Update, context:ContextTypes.DEFAULT_TYPE):
    buttons=[ [InlineKeyboardButton("Add a book",callback_data="add_book")],[InlineKeyboardButton("Delete a book",callback_data="delete_book")],[InlineKeyboardButton("View the reading list",callback_data="view_reading_list")]]
    reply_markup=InlineKeyboardMarkup(buttons)
    await update.message.reply_text("Choose an action: ",reply_markup=reply_markup)

def update_reading_list_docx(reading_list, docx_file):
    doc=Document()
    doc.add_heading("READING LIST",level=0)
    if not reading_list:
        doc.add_paragraph("The reading list is empty.")
    else:
        for book in reading_list:
            doc.add_paragraph(book)
    doc.save(docx_file)
    print("Document saved")

reading_list=[]
docx_file="reading_list.docx"

async def button_handler(update:Update, context:ContextTypes.DEFAULT_TYPE):
    query=update.callback_query
    await query.answer()
    
    reading_list=context.user_data.get("reading_list",[])
    if query.data=="add_book":
        book_name=context.user_data.get("book_name","No name")
        if book_name!="No name":
            reading_list.append(book_name)
            context.user_data["reading_list"]=reading_list
            update_reading_list_docx(reading_list, docx_file)
            await query.edit_message_text(text=f"{book_name} has been added to your reading_list")
        else:
            await query.edit_message_text(text="You didn't tell me the book name. Try again")
    elif query.data=="delete_book":
        book_name=context.user_data.get("book_name","No name")
        if book_name in reading_list:
            reading_list.remove(book_name)
            context.user_data["reading_list"]=reading_list
            update_reading_list_docx(reading_list, docx_file)
            await query.edit_message_text(text=f"{book_name} has been removed from your reading list")
        else:
            await query.edit_message_text(text=f"{book_name} was not found in your reading list")
    elif query.data=="view_reading_list":
        if os.path.exists(docx_file):
            await query.message.reply_document(document=open(docx_file,'rb' ),filename="reading_list.docx")
        else:
            await query.message.reply_text("The reading list file does not exist.")

#help command
async def help_command(update: Update, contxext:ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("/start - Get started talking to Yuki \n"
               "/help - show this help list\n" 
               "/book - Enter a genre you are interested in and get a list of books in CSV format\n"
               "/preview - Enter the name of a book and get the preview link for that book\n"
               "/list - Enter a book name to add to the reading list \n"
               "/reading_list - Used to modify and/or view your reading list\n")

if __name__=="__main__":
    print("Starting bot..........................")
    app=Application.builder().token(TELEGRAM_BOT_KEY).build()

    app.add_handler(CommandHandler("start",start_command))
    app.add_handler(CommandHandler("book",book_command))
    app.add_handler(CommandHandler("preview",preview_command))
    app.add_handler(CommandHandler("list",list_command))
    app.add_handler(CommandHandler("help",help_command))
    app.add_handler(CommandHandler("reading_list",reading_list_command))
    app.add_handler(CallbackQueryHandler(button_handler))

    print("Polling........................")
    app.run_polling(poll_interval=3)
